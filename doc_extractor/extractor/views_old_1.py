# Django Libraries
from django.shortcuts import render
from django.http import HttpResponse , JsonResponse

# Custom Libraries
import pandas as pd
import numpy as np
import joblib
from laserembeddings import Laser
from sklearn.neural_network import MLPClassifier                            # works great -- neural network
from langid import classify
from langdetect import detect
import os
import re
import pdfplumber
from docx2json import convert
import json
from docx import Document
import smbclient
import mammoth
from bs4 import BeautifulSoup

# Rest framework import
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.decorators import api_view,permission_classes,authentication_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.authentication import SessionAuthentication, BasicAuthentication , TokenAuthentication

# Other file import
from environment import MODE

if MODE == 'local':
    from .local_constants import *
else:
    from .dev_constants import *

categories = ['nutrition','ingredients','allergen statement','shelf_life_statement',
              'storage instruction','address',
              # 'warning statement',
              "gtin_number","serial_number","lot_number","expiry_date",'form_content',
              'usage instruction','pc_number','general classification',"eu_number"]

msd_categories = ['name','active_substance','excipients','form_content','method_route','warning','expiry_date',
                  'storage_instructions','precautions','marketing_company','unique_identifier','classification',
                  'usage_instruction','braille_info','mfg_date','manufacturer','packing_site','appearance',
                  'product_info','label_dosage','box_info']



# Initialize Laser
laser = Laser(path_to_bpe_codes,path_to_bpe_vocab,path_to_encoder)

# @authentication_classes([SessionAuthentication, BasicAuthentication])
@api_view()
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def extractor(request):
    content = {'message': 'Hello, World!'}
    return Response(content)
    # return render(request,'extractor/index.html')

# Create your views here.
def model_training():
    df = pd.read_excel(input_excel)
    df = df.sample(frac=1)
    X_train_laser = laser.embed_sentences(df['text'], lang='en')
    # mlp = MLPClassifier(hidden_layer_sizes=(125,), solver='adam', activation='tanh', random_state=0, shuffle=True)
    mlp = MLPClassifier(hidden_layer_sizes=(70,),solver='adam',max_iter=500,activation='tanh',random_state=0,shuffle=True)
    mlp.fit(X_train_laser, df['category'])
    joblib.dump(mlp,model_location)
    return mlp

def classifier(request):
    text = request.GET.get('text','')
    if text:
        pass
    else:
        return render(request, 'extractor/index_classifier.html')
    model = None
    if os.path.exists(model_location):
        model = joblib.load(model_location)
    else:
        model = model_training()
    # lang_detected = detect(text)
    lang_detected = classify(text)[0]
    # print('lang----->',lang_detected)
    # print(text)
    prediction = model.predict(laser.embed_sentences([text],lang=lang_detected))
    probability = model.predict_proba(laser.embed_sentences([text],lang=lang_detected))
    probability[0].sort()
    max_probability = max(probability[0])
    if (max_probability-0.35) > probability[0][-2]:
        pred_output = prediction[0]
    else:
        pred_output = 'None'
    # print(probability)
    print('{}-------------->{}'.format(max(probability[0]),pred_output))
    result = {'probability':max(probability[0]),'output':pred_output,'actual_output':prediction[0],'text':text}
    # return HttpResponse(pred_output)
    # return render(request,'extractor/doc_result.html',{'result':dict})
    return render(request,'extractor/index_result.html',result)

def prediction(text):
    model = None
    if os.path.exists(model_location):
        model = joblib.load(model_location)
    else:
        model = model_training()
    # lang_detected = detect(text)
    lang_detected = classify(text)[0]
    # print('lang----->',lang_detected)
    print(text)
    prediction = model.predict(laser.embed_sentences([text],lang=lang_detected))
    probability = model.predict_proba(laser.embed_sentences([text],lang=lang_detected))
    probability[0].sort()
    max_probability = max(probability[0])
    # if (max_probability-0.35) > probability[0][-2]:
    if max_probability > 0.63:
        pred_output = prediction[0]
    else:
        pred_output = 'None'
    print('{}-------------->{}'.format(max(probability[0]),pred_output))
    return ({'probability': max(probability[0]), 'output': pred_output, 'actual_output': prediction[0]})

def doc_extractor(request):
    final = {}
    file_name = request.GET.get('file','no file')
    if file_name == 'no file':
        return render(request, 'extractor/index.html')
    else:
        pass
    file = document_location+file_name
    doc_format = os.path.splitext(file_name)[1].lower()
    if doc_format == ".pdf":
        if os.path.exists(file):
            pdf = pdfplumber.open(file)
        else:
            return HttpResponse('File not found')
        no_of_pages = len(pdf.pages)
        tables = len(pdf.pages[0].extract_tables())
        if tables > 2:
            print('type 1 --- tables')
            for page_no in range(no_of_pages):
                page = pdf.pages[page_no]
                extracted_table = page.extract_tables()
                text = [" ".join(list(filter(None, content))).replace('\n', ' ') for table in extracted_table for content in table]
                for sentence in text:
                    unique_identifiers = Regex_parsers(sentence)
                    if unique_identifiers:
                        final = {**final, **unique_identifiers}
                    else:
                        pass
                    result = prediction(sentence)['output']
                    if result != 'None':
                        if result in final.keys():
                            final[result].append(sentence)
                        else:
                            final[result] = [sentence]
                    else:
                        pass
            if len(final['Nutrition']) > 1:
                final['Nutrition'] = final['Nutrition'][:-1]
            else:
                pass
            extracted_categories = {key:val for key, val in final.items() if key.lower() in categories}
            # return JsonResponse(extracted_categories)
            return render(request, 'extractor/doc_result.html', {'result': extracted_categories})
        else:
            print('type-2-paragraph')
            for page_no in range(no_of_pages):
                page = pdf.pages[page_no]
                extracted_text = page.extract_text()
                text = sentence_tokennizer(extracted_text)
                for sentence in text:
                    unique_identifiers = Regex_parsers(sentence)
                    if unique_identifiers:
                        final = {**final, **unique_identifiers}
                    else:
                        pass
                    result = prediction(sentence)['output']
                    if result in final.keys():
                        final[result].append(sentence)
                    else:
                        final[result] = [sentence]
            extracted_categories = {key:val for key, val in final.items() if key.lower() in categories}
            return render(request, 'extractor/doc_result.html', {'result': extracted_categories})
    elif (doc_format == '.docx') or (doc_format == '.doc'):
        doc = convert(file,sepBold=True)
        doc_to_json = json.loads(doc)
        text = doc_to_json['nonbold']
        if text:
            pass
        else:
            text = doc_to_json['text']
        for sentence in text:
            unique_identifiers = Regex_parsers(sentence)
            if unique_identifiers:
                final = {**final,**unique_identifiers}
            else:
                pass
            result = prediction(sentence)['output']
            if result in final.keys():
                final[result].append(sentence)
            else:
                final[result] = [sentence]
        # print(final)
        extracted_categories = {key: val for key, val in final.items() if key.lower() in categories}
        return render(request, 'extractor/doc_result.html', {'result': extracted_categories})
    else:
        return HttpResponse('This file format not supported currently')

def sentence_tokennizer(text):
    #sentences = re.split(r"[.!?]", text)
    # sentences = re.split(r"\.\s\n", text)
    segments = re.split(r"\n\s\n", text)
    sentences = [re.split(r"\.\s\n", seg) for seg in segments]
    # token = [re.sub(r"\d\-.*",'number',text) for sublist in sentences for text in sublist]
    token = [text for sublist in sentences for text in sublist]
    # sentences = [sent.strip() for sent in sentences if sent]
    return token

def Regex_parsers(text):
    unique_number = {}
    for key , pattern in regex_patterns.items():
        finding = re.findall(pattern,text,(re.IGNORECASE|re.MULTILINE))
        try:
            finding = str(finding[0]).strip()
        except:
            pass
        if finding:
            # print("---------************{}".format(finding))
            unique_number[key] = [finding]
        else:
            pass
    return unique_number

def msd_data_extractor(list,regex_heading_msd):
    tmp = []
    final = {}
    key = ''
    for i in range(len(list)):
        text = str(list[i])
        if re.findall(regex_heading_msd, text):
            try:
                if key != '':
                    final[key] = '\n'.join(tmp)
                else:
                    pass
                key = text
                tmp.clear()
            except:
                pass
        else:
            if i == len(list) - 1:
                tmp.append(text)
                final[key] = ' '.join(tmp)
            else:
                tmp.append(text)
    return final


def msd_prediction(text):
    model = None
    if os.path.exists(msd_model_location):
        model = joblib.load(msd_model_location)
    else:
        model = msd_model_training()
        print('new model trained')
    # lang_detected = detect(text)
    lang_detected = classify(text)[0]
    # print('lang----->',lang_detected)
    # print(text)
    prediction = model.predict(laser.embed_sentences([text],lang=lang_detected))
    probability = model.predict_proba(laser.embed_sentences([text],lang=lang_detected))
    probability[0].sort()
    max_probability = max(probability[0])
    # if (max_probability-(max_probability/2)) > probability[0][-2]:
    if max_probability > 0.60:
        pred_output = prediction[0]
    else:
        pred_output = 'None'
    print('{}-------------->{}'.format(max(probability[0]),pred_output))
    return ({'probability': max(probability[0]), 'output': pred_output, 'actual_output': prediction[0]})

def msd_model_training():
    df = pd.read_excel(msd_input_excel)
    df = df.sample(frac=1)
    X_train_laser = laser.embed_sentences(df['text'], lang='en')
    # mlp = MLPClassifier(hidden_layer_sizes=(125,), solver='adam', activation='tanh', random_state=0, shuffle=True)
    mlp = MLPClassifier(hidden_layer_sizes=(70,),solver='adam',max_iter=500,activation='tanh',random_state=0,shuffle=True)
    # mlp = MLPClassifier(hidden_layer_sizes=(70,),solver='adam',max_iter=300,activation='relu',learning_rate='constant',learning_rate_init=0.001,random_state=0,shuffle=True)
    mlp.fit(X_train_laser, df['category'])
    joblib.dump(mlp,msd_model_location)
    return mlp

# @api_view()
# @permission_classes([IsAuthenticated])
# @authentication_classes([TokenAuthentication])
def msd(request):
    final_json = {}
    # print(request.args.getlist('file'))
    # print(request.GET['file'])
    file_name_list = request.GET.getlist('file','no file')
    print('file_list',file_name_list)

    # getting value from query string
    # file_name_list = request.GET.get('file','no file')
    # print(file_name_list)
    if file_name_list == 'no file':
        return render(request, 'extractor/index_msd.html')
        # return Response({'status':'0'})
    else:
        pass
    for index , file_name in enumerate(file_name_list):
        extracted_categories = {}
        final = {}
        cate_tmp = {}
        lang_final = []
        # Reading file from storage
        if MODE == 'local':
            file = document_location+file_name
        else:
            file = get_file_smb(r"{}".format(file_name))

        doc_format = os.path.splitext(file_name)[1].lower()
        if doc_format == '.docx':
            '''
            doc = Document(file)
            try:
                file.close()
            except:
                pass
            for para in doc.paragraphs:
                p = str(para.text).strip()
                if p:
                    list.append(p)
                else:
                    pass
            extracted = msd_data_extractor(list,regex_heading_msd)
            '''
            extracted , lang = text_extraction(file)
            for key,value in extracted.items():
                if "".join(value).strip() != '':
                    result = msd_prediction(key)['output']
                    if result != 'None':
                        if result in final.keys():
                            # final[result].append(re.sub(r'\\n',' ',value).strip())
                            final[result].append(value.replace('\n',' ').strip())
                            # final[result].append({lang:value.replace('\n',' ').strip()})
                        else:
                            # final[result] = [re.sub(r'\\n',' ',value).strip()]
                            final[result] = [value.replace("\n",' ').strip()]
                            # final[result] = [{lang:value.replace('\n',' ').strip()}]
                    else:
                        pass
            # extracted_categories = {str(key): val for key, val in final.items() for lang , value in val.items() if ''.join(value).strip() != ''}
            unique = {}
            if 'unique_identifier' in final:
                unique = Regex_parsers(str(final['unique_identifier']))
                final.pop('unique_identifier')
            else:
                pass

            for cate , value in final.items():
                if cate in msd_categories_lang:
                    # if len(value) > 1:
                    if isinstance(value,list):
                        for t in value:
                            if '$$' in t:
                                list_text = t.split('$$')
                                # print('list-------->', list_text)
                                topic = ''
                                for text in list_text:
                                    try:
                                        text = text.replace('$$', ' ')
                                    except:
                                        pass
                                    if len(str(text).split()) > 2:
                                        if topic:
                                            text = topic + " " + text
                                            topic = ''
                                        lang = detect(text)
                                        if lang not in lang_final:
                                            lang_final.append(lang)
                                        else:
                                            pass
                                        if cate in cate_tmp:
                                            cate_tmp[cate].append({lang: text})
                                        else:
                                            cate_tmp[cate] = [{lang: text}]
                                    else:
                                        topic = text
                                        # print('topic----->',topic)
                            else:
                                lang = detect(t)
                                cate_tmp[cate] = [{lang: t}]
                    else:
                        pass
                elif '$$' in value[0]:
                    cate_tmp[cate] = [value[0].replace('$$',' ')]
                else:
                    cate_tmp[cate] = [value[0]]
                    continue

            print(final)
            status = {'status':'1','language': lang_final,'file_name':[file_name]}
            extracted_categories = {**status,**cate_tmp,**unique}
            final_json[index] = extracted_categories

            # return JsonResponse(extracted_categories)
            # return Response(extracted_categories)
            # return render(request, 'extractor/doc_result.html', {'result': extracted_categories})
        else:
            status = {'status': '0','file_name': [file_name]}
            final_json[index] = status
            # return JsonResponse(status)
            # return Response(status)
    return JsonResponse(final_json)


def get_file_smb(file_name):
    data = ''
    try:
        data = smbclient.open_file(r"{}".format(file_name),mode='rb',username=smb_username,password=smb_password)
        print('file found')
    except:
        smbclient.reset_connection_cache()
        data = smbclient.open_file(r"{}".format(file_name), mode='rb',username=smb_username,password=smb_password)
    finally:
        return data

def text_extraction(file,method=None):
    tmp = []
    final = {}
    key = ''
    lang = []
    html = mammoth.convert_to_html(file).value
    try:
        file.close()
        html.close()
    except:
        pass

    '''
    soup = BeautifulSoup(html,'html.parser')
    paragraphs = soup.find_all('p')
    # list = [ele.text for ele in paragraphs]
    list = [ele.next for ele in paragraphs]
    '''
    soup = BeautifulSoup(html, 'html.parser')
    paragraphs = soup.find_all('p')
    soup_list = [ele.text for ele in paragraphs]
    # ------
    paragraph = html.split('</p>')
    list = [string.replace('<p>', "") for string in paragraph if string]
    # -----
    for i in range(len(list)):
        text = str(list[i])
        if re.findall(regex_heading_msd, text):
            try:
                if key != '':
                    final[key] = '$$'.join(tmp)
                else:
                    pass
                key = re.sub(r'<.*?>','',soup_list[i])
                # print(key)
                tmp.clear()
            except:
                pass
        else:
            if i == len(list) - 1:
                text = text.replace('<strong>','<b>').replace('</strong>','</b>')
                tmp.append(text)
                # if len(text.split()) > 6:
                #     lang.append(classify(text)[0])
                # else:
                #     pass
                final[key] = '$$'.join(tmp)
            else:
                text = text.replace('<strong>', '<b>').replace('</strong>', '</b>')
                # if len(text.split()) > 6:
                #     lang.append(classify(text)[0])
                # else:
                #     pass
                tmp.append(text)
    # return final , max(lang,key=lang.count)
    return final , lang


#
# x = get_file_smb(r"\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\Book1.xlsx")
#
# x = get_file_smb(r"\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\QRD.docx")
#
# x = get_file_smb(r"\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\Variant Design Creation.xlsx")
#
# xx = x.read()

# original = "\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\Book1.xlsx"
#
# x = get_file_smb(r"\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\Book1.xlsx")
#
# import smbclient

# data = smbclient.open_file(r"\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\QRD.docx",mode='rb',username="weblogic",password="417@sia123")

# data.close()

# x = data
#
# x.close()
#
# data.read()


# data.close()
#
# xx = x.read()

# data1 = smbclient.open_file(r"\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\Variant Design Creation.xlsx",mode='rb',username="weblogic",password="417@sia123")
#
# data1.close()

# # data = smbclient.open_file("\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\/Variant Design Creation.xlsx",mode='r',username="weblogic",password="417@sia123",encoding='utf8')
# data = smbclient.open_file("\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\/401030005.xml",mode='r',username="weblogic",password="417@sia123")
# #
# smbclient.reset_connection_cache()
#
# smbclient.register_session()
#
# file = data.read()
# #
# # data1 = data.decode('utf8')
# #
# # import io
# #
# # encoding='utf-8-sig'
# import pandas as pd
# with smbclient.open_file(r'\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\/Book1.xlsx',mode='rb',username="weblogic",password="417@sia123",encoding='latin') as f:
#     t = f.read()
#
# z = smbclient.open_file(r'\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\/Book1.xlsx',mode='rb',username="weblogic",password="417@sia123")
#
# destination = r"/Users/VIJAYKANAGARAJ/PycharmProjects/Schawk_document_xml/Dataset/sssssss.xlsx"
#
# x = smbclient.(r'\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\/Book1.xlsx',mode='r+',username="weblogic",password="417@sia123",encoding='latin',dst=destination)
#
#
# smb = smbclient.SambaClient()
#
#     t = open(f,'r')
# #
# # import chardet
# #
# # enc = chardet.detect(f)
#
# x = get_file_smb(r"\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\Book1.xlsx")
#
# with smbclient.open_file(r'\\chenfile05.asia.schawk.com\Tornado\TORNADO_TESTING\020_Supplied\A01\OLD\Book1.xlsx', 'r') as f:
#     x  = f
#
# import tempfile
# from smb.SMBConnection import SMBConnection
#
# conn = SMBConnection('weblogic', '417@sia123',"chenfile05","chenfile05.asia.schawk.com", use_ntlm_v2 = True)
# assert conn.connect("//chenfile05.asia.schawk.com", 445)
#
#
# inp_e = r"/Users/VIJAYKANAGARAJ/PycharmProjects/Schawk_document_xml/Dataset/QRD_doc.docx"
#
#
# d = open(inp_e,mode='rb')
#
# from docx import Document
# #
# doc = Document(data)
#
# data.close()
# #
# for para in doc.paragraphs:
#     print(para.text)
# #
# df = pd.read_excel(z)      #
